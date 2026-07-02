"""
api.py — FastAPI backend for the RFP web UI

Endpoints:
  POST /upload          — upload + parse an RFP document
  GET  /documents       — list all uploaded documents
  GET  /responses/{id}  — get generated responses for a document
  POST /regenerate      — regenerate a single answer
  GET  /knowledge       — knowledge base stats
  POST /export          — export responses to xlsx/docx/pdf

Run with:
  uvicorn api:app --reload --port 8000
"""

import os
import io
import json
import uuid
import time
import threading
import zipfile as _zipfile

UPLOAD_ALLOWED_EXTS = {".xlsx", ".xls", ".csv"}
MAX_UPLOAD_BYTES = 20 * 1024 * 1024

# Magic bytes for real content-type detection (no system library needed)
_XLSX_MAGIC = b"PK\x03\x04"          # ZIP archive (xlsx is a zip)
_XLS_MAGIC  = b"\xd0\xcf\x11\xe0"    # OLE2 compound document
_CSV_CHARS  = set(b"abcdefghijklmnopqrstuvwxyzABCDEFGHIJKLMNOPQRSTUVWXYZ0123456789,;|\t\r\n \"'-_.:/()&+@#$ ")

def _validate_upload(filename: str, content: bytes) -> None:
    if len(content) > MAX_UPLOAD_BYTES:
        raise HTTPException(400, f"File too large — max {MAX_UPLOAD_BYTES // 1048576} MB")
    ext = Path(filename or "").suffix.lower()
    if ext not in UPLOAD_ALLOWED_EXTS:
        raise HTTPException(400, f"Unsupported extension '{ext}'. Allowed: {sorted(UPLOAD_ALLOWED_EXTS)}")
    # Verify actual content matches claimed extension
    head = content[:4]
    if ext == ".xlsx":
        if head != _XLSX_MAGIC:
            raise HTTPException(400, "File content is not a valid .xlsx spreadsheet")
        try:
            _zipfile.ZipFile(io.BytesIO(content))
        except _zipfile.BadZipFile:
            raise HTTPException(400, "File is corrupt or not a valid .xlsx")
    elif ext == ".xls":
        if head != _XLS_MAGIC:
            raise HTTPException(400, "File content is not a valid .xls spreadsheet")
    elif ext == ".csv":
        sample = content[:1024]
        if not all(b in _CSV_CHARS for b in sample):
            raise HTTPException(400, "File does not appear to be valid CSV text")
from datetime import datetime, timezone
from pathlib import Path
from typing import Optional

from fastapi import FastAPI, UploadFile, File, Form, HTTPException, BackgroundTasks, Request
from slowapi import Limiter, _rate_limit_exceeded_handler
from slowapi.util import get_remote_address
from slowapi.errors import RateLimitExceeded
from fastapi import Depends
from database import create_db_and_tables, engine
from sqlmodel import Session
from models import OriginalDocument
from auth import require, current_user, User, list_users, upsert_user, delete_user, ROLES
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import StreamingResponse
from pydantic import BaseModel

from parser import parse_rfp_document, items_to_dict
from retriever import ask
from confidence import compute_confidence

# Max questions to auto-answer in one upload (rest can be done via /generate/{doc_id})
MAX_AUTO_ANSWER = 20
# Delay between Groq calls to avoid rate limiting (seconds)
GROQ_DELAY = 0.3

limiter = Limiter(key_func=get_remote_address)
app = FastAPI(title="RFP Pilot API", version="2.0")

from review_workflow import router as review_router
app.include_router(review_router)

@app.on_event("startup")
async def startup_event():
    """Create DB tables on boot (idempotent — safe to run every startup)."""
    create_db_and_tables()
app.state.limiter = limiter
app.add_exception_handler(RateLimitExceeded, _rate_limit_exceeded_handler)

@app.on_event("startup")
def _startup_checks():
    if os.getenv("AUTH_DISABLED", "").lower() == "true":
        import sys
        env = os.getenv("ENV","development").lower()
        sep = "!"*70
        print(f"\n{sep}\n  WARNING: AUTH_DISABLED=true — ALL REQUESTS RUN AS ADMIN", file=sys.stderr)
        if env not in ("development","dev","local"):
            print("  NON-DEVELOPMENT env detected. Stop this service immediately!", file=sys.stderr)
        print(sep, file=sys.stderr)


# Lock CORS down by default; override per environment with e.g.
#   ALLOWED_ORIGINS="https://rfi.yourcompany.com,https://staging.yourcompany.com"
# Set ALLOWED_ORIGINS="*" only for local experimentation.
_origins = [o.strip() for o in os.getenv(
    "ALLOWED_ORIGINS",
    "http://localhost:3000,http://127.0.0.1:3000",
).split(",") if o.strip()]

app.add_middleware(
    CORSMiddleware,
    allow_origins=_origins,
    allow_methods=["*"],
    allow_headers=["*"],
)

# In-memory store (swap for a real DB in production)
DOCUMENTS: dict = {}    # doc_id → {meta, items, responses}

def _require_doc_access(doc_id: str, user: "User"):
    """H3: prevent IDOR — only the owner or an admin may touch a document."""
    doc = DOCUMENTS.get(doc_id)
    if not doc:
        raise HTTPException(404, "Document not found")
    if user.role != "admin" and doc.get("owner_email") not in (None, user.email):
        raise HTTPException(403, "You do not have access to this document")
    return doc


def _persist_original_document(doc_id: str, filename: str, content: bytes, uploaded_by: str) -> None:
    """
    Phase 5: store the raw uploaded workbook keyed by doc_id so the review
    export endpoint can later reopen it and write approved answers back into
    the original sheet structure. Called after successful parsing so we don't
    keep orphan bytes for files that failed to parse.
    """
    ext = Path(filename or "").suffix.lower().lstrip(".") or "bin"
    with Session(engine) as s:
        s.add(OriginalDocument(
            doc_id=doc_id,
            filename=filename or f"upload.{ext}",
            content_type=ext,
            content=content,
            uploaded_by=uploaded_by,
        ))
        s.commit()


# ── Health ────────────────────────────────────────────────────

@app.get("/health")
def health():
    """Cheap liveness probe for the frontend status indicator and load balancers."""
    return {"status": "ok", "version": app.version}


# ── RFI history (org-wide, server-side) ──────────────────────
# Replaces per-browser localStorage history: every user with access sees
# the same ledger of processed RFIs. JSON-file storage (swap for a DB
# table later, same as users.json).

HISTORY_FILE = Path(os.getenv("HISTORY_FILE", str(Path(__file__).parent / "history.json")))
_hist_lock = threading.Lock()


def _load_history() -> list[dict]:
    if HISTORY_FILE.exists():
        try:
            return json.loads(HISTORY_FILE.read_text())
        except Exception:
            return []
    return []


def _save_history(entries: list[dict]) -> None:
    HISTORY_FILE.write_text(json.dumps(entries, indent=2))
    try: HISTORY_FILE.chmod(0o600)
    except OSError: pass


@app.get("/history", dependencies=[Depends(require("kb_read"))])
def get_history():
    from sqlmodel import Session as _Session, select as _select
    from database import engine as _engine
    from models import HistoryEntry as _HistoryEntry
    with _Session(_engine) as _s:
        rows = _s.exec(_select(_HistoryEntry)).all()
    entries = [r.payload for r in rows if r.payload]
    entries.sort(key=lambda e: e.get("uploadedAt", ""), reverse=True)
    return {"entries": entries[:200]}


@app.put("/history")
def upsert_history(entry: dict, user: User = Depends(require("generate"))):
    """Create or merge-update an RFI history entry by id (Postgres-backed)."""
    from sqlmodel import Session as _Session
    from database import engine as _engine
    from models import HistoryEntry as _HistoryEntry
    from sqlalchemy.orm.attributes import flag_modified
    eid = (entry.get("id") or "").strip()
    if not eid:
        raise HTTPException(400, "entry.id is required")
    with _Session(_engine) as _s:
        row = _s.get(_HistoryEntry, eid)
        if row is not None:
            merged = dict(row.payload or {})
            merged.update(entry)
            merged["updatedAt"] = datetime.now(timezone.utc).isoformat()
            row.payload = merged
            row.owner = merged.get("owner", row.owner)
            row.filename = merged.get("filename", merged.get("name", row.filename))
            flag_modified(row, "payload")
            _s.commit()
            _s.refresh(row)
            return row.payload
        entry["owner"] = user.email
        entry["ownerName"] = user.name
        entry["updatedAt"] = datetime.now(timezone.utc).isoformat()
        _s.add(_HistoryEntry(
            id=eid,
            owner=user.email,
            filename=entry.get("filename", entry.get("name", "unknown")),
            row_count=entry.get("rowCount", entry.get("row_count", 0)),
            payload=entry,
        ))
        _s.commit()
        return entry


@app.delete("/history/{entry_id}")
def delete_history(entry_id: str, user: User = Depends(current_user)):
    from sqlmodel import Session as _Session
    from database import engine as _engine
    from models import HistoryEntry as _HistoryEntry
    with _Session(_engine) as _s:
        row = _s.get(_HistoryEntry, entry_id)
        if row is None:
            return {"deleted": entry_id}
        if user.role != "admin" and row.owner != user.email:
            raise HTTPException(403, "Only the owner or an admin can remove a history entry")
        _s.delete(row)
        _s.commit()
    return {"deleted": entry_id}


# ── Review queue: server-side persistence (M4) ───────────────
REVIEW_FILE = Path(os.getenv("REVIEW_FILE", str(Path(os.getenv("FEEDBACK_LOG", str(Path(__file__).parent / "feedback_log.jsonl"))).parent / "review_queue.json")))
_review_lock = threading.Lock()


def _load_review() -> list[dict]:
    if REVIEW_FILE.exists():
        try:
            return json.loads(REVIEW_FILE.read_text())
        except Exception:
            return []
    return []


def _save_review(items: list[dict]) -> None:
    REVIEW_FILE.parent.mkdir(parents=True, exist_ok=True)
    REVIEW_FILE.write_text(json.dumps(items, indent=2))
    try: REVIEW_FILE.chmod(0o600)
    except OSError: pass


@app.get("/review-queue", dependencies=[Depends(require("kb_read"))])
def get_review_queue():
    return {"responses": _load_review()}


@app.put("/review-queue")
def put_review_queue(payload: dict, user: User = Depends(require("generate"))):
    responses = payload.get("responses")
    if not isinstance(responses, list):
        raise HTTPException(400, "responses must be a list")
    with _review_lock:
        _save_review(responses[-2000:])
    return {"saved": len(responses[-2000:])}


@app.get("/stats", dependencies=[Depends(require("kb_read"))])
def get_stats():
    items = _load_review()
    def _count(status): return sum(1 for r in items if r.get("status") == status)
    scores = [r.get("confidence", {}).get("score", 0) for r in items
              if r.get("status") not in ("generating", "error")]
    scores = [s for s in scores if s]
    return {
        "total":         len(items),
        "generated":     _count("generated"),
        "needsReview":   _count("needs_review"),
        "approved":      _count("approved"),
        "rejected":      _count("rejected"),
        "exported":      _count("exported"),
        "lowConfidence": sum(1 for r in items if (r.get("confidence", {}).get("score", 0) or 0) < 0.7),
        "avgConfidence": (sum(scores) / len(scores)) if scores else 0,
    }


# ── Knowledge base: Drive sync ───────────────────────────────

SYNC_STATE = {"running": False, "last_started": None}

@app.post("/knowledge/sync", dependencies=[Depends(require("kb_write"))])
def trigger_drive_sync(background: BackgroundTasks):
    """Run the Google Drive ingestion (ingest.run_ingestion) in the
    background. The Sync Drive button previously posted to a nonexistent
    endpoint — this makes it real."""
    if SYNC_STATE["running"]:
        return {"status": "already_running", "started_at": SYNC_STATE["last_started"]}

    def _run():
        SYNC_STATE["running"] = True
        SYNC_STATE["last_started"] = datetime.now(timezone.utc).isoformat()
        try:
            from ingest import run_ingestion
            run_ingestion()
        except Exception as e:
            print(f"  [sync] Drive ingestion failed: {e}")
        finally:
            SYNC_STATE["running"] = False

    background.add_task(_run)
    return {"status": "started"}


@app.get("/knowledge/sync", dependencies=[Depends(require("kb_read"))])
def sync_status():
    return SYNC_STATE


# ── Identity & user management ───────────────────────────────

@app.get("/me")
def me(user: User = Depends(current_user)):
    """Who am I and what can I do — drives all frontend gating."""
    return {
        "email": user.email,
        "name": user.name,
        "role": user.role,
        "capabilities": sorted(user.capabilities),
    }


class UserUpsert(BaseModel):
    email: str
    role: str


@app.get("/admin/users")
def admin_list_users(user: User = Depends(require("manage_users"))):
    return {"users": list_users(), "roles": list(ROLES)}


@app.put("/admin/users")
def admin_upsert_user(body: UserUpsert, user: User = Depends(require("manage_users"))):
    return upsert_user(body.email, body.role, added_by=user.email)


@app.delete("/admin/users/{email}")
def admin_delete_user(email: str, user: User = Depends(require("manage_users"))):
    if email.strip().lower() == user.email:
        raise HTTPException(400, "You cannot remove your own account")
    delete_user(email)
    return {"deleted": email}


# ── Single-question answer (used by the column-mapper UI) ────

VALID_AVAILABILITY = ("Yes", "No", "Partial", "Unknown")

def normalise_availability(raw: str) -> str:
    """
    The LLM occasionally returns variants like "yes", "YES — fully supported",
    or "Partially". The frontend's AvailabilityLabel type only allows
    Yes / No / Partial / Unknown, so coerce everything into that taxonomy
    instead of leaking free-form strings into the UI and exports.
    """
    cleaned = (raw or "").strip().lower()
    if cleaned.startswith("yes"):
        return "Yes"
    if cleaned.startswith("no"):
        return "No"
    if cleaned.startswith("partial"):
        return "Partial"
    return "Unknown"


import secrets as _secrets
def _safe_error(e: Exception, ref: str = "") -> HTTPException:
    ref = ref or _secrets.token_hex(4)
    print(f"  [error:{ref}] {type(e).__name__}: {e}")
    return HTTPException(500, detail={"message": "An internal error occurred.", "ref": ref})

def is_rate_limit_error(err: Exception) -> bool:
    msg = str(err)
    return "429" in msg or "rate_limit" in msg.lower() or "rate limit" in msg.lower()


class QuestionRequest(BaseModel):
    question: str
    section: str = ""


# Guarded with "ask" (not "generate") because the Assistant page shares
# this endpoint and read-only users may ask questions. Bulk RFI flows
# (parse / regenerate / upload / export / corrections) remain role-gated.
@app.post("/answer", dependencies=[Depends(require("ask"))])
@limiter.limit("20/minute")
def answer_question(request: Request, req: QuestionRequest, user: User = Depends(current_user)):
    """
    Answer a single question from the knowledge base.
    Used by the frontend column-mapper UI — the user picks their own
    question column, selects rows, and we answer each one individually.
    Much faster than /upload for interactive use.
    """
    if not req.question or len(req.question.strip()) < 5:
        raise HTTPException(400, "Question is too short.")

    try:
        result = ask(req.question.strip())
    except Exception as e:
        # Surface provider throttling as a real 429 so the frontend can
        # pause the batch and offer Resume, instead of treating it as a
        # generic server error.
        if is_rate_limit_error(e):
            raise HTTPException(429, "Model provider rate limit reached. Please retry shortly.")
        raise _safe_error(e)

    # Parse availability + remarks out of the raw answer string
    availability = ""
    remarks = ""
    for line in result["answer"].split("\n"):
        line = line.strip()
        if line.startswith("AVAILABILITY:"):
            availability = line.replace("AVAILABILITY:", "").strip()
        elif line.startswith("REMARKS:"):
            remarks = line.replace("REMARKS:", "").strip()

    return {
        "question":      req.question,
        "section":       req.section,
        "availability":  normalise_availability(availability),
        "remarks":       remarks or result["answer"],
        "sources":       result["sources"],
        "confidence":    result["confidence"],
    }



# ── Upload & parse ────────────────────────────────────────────

@app.post("/documents/upload", dependencies=[Depends(require("generate"))])
async def upload_document_raw(
    file: UploadFile = File(...),
    doc_id: str = Form(...),
    user: User = Depends(current_user),
):
    """
    Phase 5: idempotent raw-bytes upload keyed by client-supplied doc_id.
    Populates original_documents for the review-workflow export endpoint
    to reopen later. Called by the frontend at submit-for-review time so
    the reviewer/submitter can export the answered file after approval.

    Returns 200 with a small JSON payload in both cases (new or existing).
    """
    content = await file.read()
    _validate_upload(file.filename or "", content)

    with Session(engine) as s:
        existing = s.get(OriginalDocument, doc_id)
        if existing:
            return {
                "doc_id": doc_id,
                "status": "already_stored",
                "bytes": len(existing.content),
                "filename": existing.filename,
            }
        ext = Path(file.filename or "").suffix.lower().lstrip(".") or "bin"
        s.add(OriginalDocument(
            doc_id=doc_id,
            filename=file.filename or f"upload.{ext}",
            content_type=ext,
            content=content,
            uploaded_by=user.email,
        ))
        s.commit()

    return {
        "doc_id": doc_id,
        "status": "stored",
        "bytes": len(content),
        "filename": file.filename or f"upload.{ext}",
    }


@app.post("/parse", dependencies=[Depends(require("generate"))])
async def parse_only(file: UploadFile = File(...), user: User = Depends(current_user)):
    """
    FAST endpoint: parse document and return extracted questions immediately.
    No answer generation — use /generate/{doc_id} after this to get answers.
    Returns in ~1 second regardless of document size.
    """
    content = await file.read()
    _validate_upload(file.filename or "", content)
    buf = io.BytesIO(content)

    try:
        items = parse_rfp_document(buf, file.filename)
    except Exception as e:
        raise _safe_error(e)

    if not items:
        raise HTTPException(422, "No questions or requirements found in the document.")

    doc_id = str(uuid.uuid4())[:8]
    # Phase 5: persist the raw workbook so the export endpoint can write
    # approved answers back into the original sheet structure later.
    _persist_original_document(doc_id, file.filename or "", content, user.email)

    DOCUMENTS[doc_id] = {
        "id":          doc_id,
        "owner_email": user.email,
        "filename":    file.filename,
        "uploaded_at": datetime.now(timezone.utc).isoformat(),
        "item_count":  len(items),
        "items":       items_to_dict(items),
        "responses":   [],
        "status":      "parsed",   # answers not yet generated
    }

    return {
        "doc_id":          doc_id,
        "filename":        file.filename,
        "questions_found": len(items),
        "status":          "parsed",
        "message":         f"Extracted {len(items)} questions. Call POST /generate/{doc_id} to generate answers.",
        "questions":       items_to_dict(items),
    }


@app.post("/upload", dependencies=[Depends(require("kb_write"))])
async def upload_document(file: UploadFile = File(...), user: User = Depends(current_user)):
    """
    Upload an RFP/RFI document.
    Parses AND generates answers for up to MAX_AUTO_ANSWER questions.
    For large documents use /parse first, then /generate/{doc_id}.
    """
    content = await file.read()
    _validate_upload(file.filename or "", content)
    buf = io.BytesIO(content)

    # 1. Parse
    try:
        items = parse_rfp_document(buf, file.filename)
    except Exception as e:
        raise _safe_error(e)

    if not items:
        raise HTTPException(422, "No questions or requirements found in the document.")

    # 2. Cap at MAX_AUTO_ANSWER to avoid timeout
    items_to_answer = items[:MAX_AUTO_ANSWER]
    skipped = len(items) - len(items_to_answer)

    # 3. Generate answers with delay between calls
    responses = []
    for item in items_to_answer:
        try:
            result = ask(item.question)
            responses.append({
                "id":         item.id,
                "section":    item.section,
                "subsection": item.subsection,
                "question":   item.question,
                "item_type":  item.item_type,
                "priority":   item.priority,
                "answer":     result["answer"],
                "sources":    result["sources"],
                "confidence": result["confidence"],
            })
        except Exception as e:
            # Don't let one failed item kill the whole batch
            responses.append({
                "id":         item.id,
                "section":    item.section,
                "subsection": item.subsection,
                "question":   item.question,
                "item_type":  item.item_type,
                "priority":   item.priority,
                "answer":     f"AVAILABILITY: Unknown\nREMARKS: Generation failed: {str(e)}",
                "sources":    [],
                "confidence": {"score": 0.0, "label": "low", "color": "red", "breakdown": {}},
            })
        time.sleep(GROQ_DELAY)

    # 4. Store all items but only answered responses
    doc_id = str(uuid.uuid4())[:8]
    # Phase 5: persist the raw workbook so the export endpoint can write
    # approved answers back into the original sheet structure later.
    _persist_original_document(doc_id, file.filename or "", content, user.email)

    DOCUMENTS[doc_id] = {
        "id":          doc_id,
        "owner_email": user.email,
        "filename":    file.filename,
        "uploaded_at": datetime.now(timezone.utc).isoformat(),
        "item_count":  len(items),
        "items":       items_to_dict(items),
        "responses":   responses,
        "status":      "partial" if skipped > 0 else "processed",
    }

    avg_conf = sum(r["confidence"]["score"] for r in responses) / len(responses) if responses else 0

    return {
        "doc_id":           doc_id,
        "filename":         file.filename,
        "questions_found":  len(items),
        "questions_answered": len(responses),
        "questions_pending": skipped,
        "avg_confidence":   round(avg_conf, 3),
        "status":           "partial" if skipped > 0 else "processed",
        "message":          f"Answered {len(responses)}/{len(items)}. Call POST /generate/{doc_id} for remaining." if skipped else None,
        "responses":        responses,
    }


@app.post("/generate/{doc_id}", dependencies=[Depends(require("generate"))])
async def generate_remaining(doc_id: str, user: User = Depends(current_user)):
    """
    Generate answers for any unanswered questions in a document.
    Call this after /parse or to continue a partial /upload.
    Answers up to MAX_AUTO_ANSWER more questions per call.
    """
    doc = _require_doc_access(doc_id, user)
    all_items = doc["items"]
    answered_ids = {r["id"] for r in doc["responses"]}
    pending = [i for i in all_items if i["id"] not in answered_ids]

    if not pending:
        return {"message": "All questions already answered.", "doc_id": doc_id}

    batch = pending[:MAX_AUTO_ANSWER]
    new_responses = []

    for item in batch:
        try:
            result = ask(item["question"])
            new_responses.append({
                "id":         item["id"],
                "section":    item["section"],
                "subsection": item["subsection"],
                "question":   item["question"],
                "item_type":  item["item_type"],
                "priority":   item["priority"],
                "answer":     result["answer"],
                "sources":    result["sources"],
                "confidence": result["confidence"],
            })
        except Exception as e:
            new_responses.append({
                "id":         item["id"],
                "section":    item["section"],
                "subsection": item["subsection"],
                "question":   item["question"],
                "item_type":  item["item_type"],
                "priority":   item["priority"],
                "answer":     f"AVAILABILITY: Unknown\nREMARKS: Generation failed: {str(e)}",
                "sources":    [],
                "confidence": {"score": 0.0, "label": "low", "color": "red", "breakdown": {}},
            })
        time.sleep(GROQ_DELAY)

    doc["responses"].extend(new_responses)
    still_pending = len(pending) - len(batch)
    if still_pending == 0:
        doc["status"] = "processed"

    return {
        "doc_id":            doc_id,
        "newly_answered":    len(new_responses),
        "total_answered":    len(doc["responses"]),
        "still_pending":     still_pending,
        "status":            doc["status"],
        "responses":         new_responses,
    }


# ── Documents list ────────────────────────────────────────────

@app.get("/documents", dependencies=[Depends(require("kb_read"))])
def list_documents():
    """
    Returns documents from Qdrant payload metadata — reflects the real knowledge base.
    Falls back to in-memory DOCUMENTS store if Qdrant unavailable.
    """
    try:
        from ingest import qdrant, COLLECTION
        # Scroll through all points to get unique source files
        seen = {}
        offset = None
        while True:
            result = qdrant.scroll(
                collection_name=COLLECTION,
                scroll_filter=None,
                limit=100,
                offset=offset,
                with_payload=True,
                with_vectors=False,
            )
            points, offset = result
            for p in points:
                src = p.payload.get("source_file", "")
                if not src:
                    continue
                if src not in seen:
                    seen[src] = {
                        "id":           p.payload.get("file_id", src),
                        "filename":     src,
                        "status":       "indexed",
                        "vectorCount":  0,
                        "uploadDate":   p.payload.get("upload_date", ""),
                        "modifiedDate": p.payload.get("upload_date", ""),
                        "source":       p.payload.get("source_type", "drive"),
                        "tags":         p.payload.get("tags", []),
                        "sizeBytes":    0,
                    }
                seen[src]["vectorCount"] += 1
            if offset is None:
                break
        return list(seen.values())
    except Exception as e:
        # Fallback to in-memory store
        return [
            {
                "id":         d["id"],
                "filename":   d["filename"],
                "uploadDate": d["uploaded_at"],
                "status":     d["status"],
                "vectorCount": 0,
                "source":     "manual",
                "tags":       [],
                "sizeBytes":  0,
            }
            for d in DOCUMENTS.values()
        ]


# ── Responses for a document ──────────────────────────────────

@app.get("/responses/{doc_id}", dependencies=[Depends(require("kb_read"))])
def get_responses(doc_id: str, user: User = Depends(current_user)):
    doc = _require_doc_access(doc_id, user)
    return {
        "doc_id":    doc_id,
        "filename":  doc["filename"],
        "responses": doc["responses"],
    }


# ── Regenerate a single answer ────────────────────────────────

class RegenerateRequest(BaseModel):
    doc_id:    str
    item_id:   str
    question:  Optional[str] = None   # override if user edited the question

@app.post("/regenerate", dependencies=[Depends(require("generate"))])
def regenerate_answer(req: RegenerateRequest, user: User = Depends(current_user)):
    _require_doc_access(req.doc_id, user)

    doc = DOCUMENTS[req.doc_id]

    # Find the item
    target = None
    for r in doc["responses"]:
        if r["id"] == req.item_id:
            target = r
            break

    if not target:
        raise HTTPException(404, f"Item {req.item_id} not found")

    question = req.question or target["question"]
    result   = ask(question)

    # Update in store
    target["answer"]     = result["answer"]
    target["sources"]    = result["sources"]
    target["confidence"] = result["confidence"]
    if req.question:
        target["question"] = req.question

    return target


# ── Knowledge base stats ──────────────────────────────────────

@app.get("/knowledge", dependencies=[Depends(require("kb_read"))])
def knowledge_stats():
    try:
        from ingest import qdrant, COLLECTION
        info = qdrant.get_collection(COLLECTION)
        vector_count = info.points_count or info.vectors_count or 0
        # Count unique source files
        seen = set()
        offset = None
        while True:
            result = qdrant.scroll(
                collection_name=COLLECTION,
                scroll_filter=None,
                limit=200,
                offset=offset,
                with_payload=["source_file"],
                with_vectors=False,
            )
            points, offset = result
            for p in points:
                src = p.payload.get("source_file","")
                if src: seen.add(src)
            if offset is None:
                break
        return {
            "vector_count":    vector_count,
            "document_count":  len(seen),
            "collection":      "rfi_knowledge_base",
            "status":          "active",
            "last_synced":     datetime.now(timezone.utc).isoformat(),
            "drive_connected": True,
        }
    except Exception:
        return {
            "vector_count":   0,
            "document_count": 0,
            "collection":     "rfi_knowledge_base",
            "status":         "offline",
        }


# ── Export ────────────────────────────────────────────────────

class ExportRequest(BaseModel):
    doc_id: str
    format: str   # "xlsx" | "docx" | "pdf"

@app.post("/export", dependencies=[Depends(require("export"))])
def export_responses(req: ExportRequest, user: User = Depends(current_user)):
    _require_doc_access(req.doc_id, user)

    doc  = DOCUMENTS[req.doc_id]
    rows = doc["responses"]

    if req.format == "xlsx":
        buf = _export_xlsx(rows, doc["filename"])
        media_type = "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"
        filename   = doc["filename"].rsplit(".", 1)[0] + "_responses.xlsx"

    elif req.format == "docx":
        buf = _export_docx(rows, doc["filename"])
        media_type = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        filename   = doc["filename"].rsplit(".", 1)[0] + "_responses.docx"

    else:
        raise HTTPException(400, "Supported formats: xlsx, docx")

    return StreamingResponse(
        buf,
        media_type=media_type,
        headers={"Content-Disposition": f'attachment; filename="{filename}"'},
    )


def _export_xlsx(rows: list, source_filename: str) -> io.BytesIO:
    import openpyxl
    from openpyxl.styles import Font, PatternFill, Alignment

    wb = openpyxl.Workbook()
    ws = wb.active
    ws.title = "RFP Responses"

    # Header
    headers = ["#", "Section", "Question", "Availability", "Response", "Confidence", "Sources"]
    for col, h in enumerate(headers, start=1):
        cell = ws.cell(row=1, column=col, value=h)
        cell.font = Font(bold=True, color="FFFFFF")
        cell.fill = PatternFill("solid", fgColor="1A2440")
        cell.alignment = Alignment(wrap_text=True)

    ws.column_dimensions["A"].width = 5
    ws.column_dimensions["B"].width = 20
    ws.column_dimensions["C"].width = 40
    ws.column_dimensions["D"].width = 14
    ws.column_dimensions["E"].width = 60
    ws.column_dimensions["F"].width = 12
    ws.column_dimensions["G"].width = 30

    # Colour fills for confidence
    green_fill  = PatternFill("solid", fgColor="0D2E1A")
    amber_fill  = PatternFill("solid", fgColor="2A1F0A")
    red_fill    = PatternFill("solid", fgColor="2A0F0F")

    for i, row in enumerate(rows, start=2):
        availability = ""
        remarks      = ""
        for line in row["answer"].split("\n"):
            if line.startswith("AVAILABILITY:"):
                availability = line.replace("AVAILABILITY:", "").strip()
            elif line.startswith("REMARKS:"):
                remarks = line.replace("REMARKS:", "").strip()

        conf_score = row["confidence"]["score"]
        conf_label = row["confidence"]["label"]
        conf_fill  = {"high": green_fill, "medium": amber_fill, "low": red_fill}.get(conf_label, amber_fill)

        ws.cell(row=i, column=1, value=row["id"])
        ws.cell(row=i, column=2, value=row["section"])
        cell_q = ws.cell(row=i, column=3, value=row["question"])
        cell_q.alignment = Alignment(wrap_text=True)
        ws.cell(row=i, column=4, value=availability)
        cell_r = ws.cell(row=i, column=5, value=remarks)
        cell_r.alignment = Alignment(wrap_text=True)
        conf_cell = ws.cell(row=i, column=6, value=conf_score)
        conf_cell.fill = conf_fill
        ws.cell(row=i, column=7, value=", ".join(row["sources"]))

        ws.row_dimensions[i].height = 60

    buf = io.BytesIO()
    wb.save(buf)
    buf.seek(0)
    return buf


def _export_docx(rows: list, source_filename: str) -> io.BytesIO:
    from docx import Document
    from docx.shared import Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    doc = Document()
    doc.add_heading(f"RFP Responses: {source_filename}", 0)
    doc.add_paragraph(f"Generated: {datetime.now().strftime('%Y-%m-%d %H:%M')}")
    doc.add_paragraph()

    current_section = ""
    for row in rows:
        if row["section"] != current_section:
            current_section = row["section"]
            doc.add_heading(current_section, level=1)

        availability = ""
        remarks      = ""
        for line in row["answer"].split("\n"):
            if line.startswith("AVAILABILITY:"):
                availability = line.replace("AVAILABILITY:", "").strip()
            elif line.startswith("REMARKS:"):
                remarks = line.replace("REMARKS:", "").strip()

        p = doc.add_paragraph()
        p.add_run(f"[{row['id']}] ").bold = True
        p.add_run(row["question"])

        p2 = doc.add_paragraph()
        p2.add_run(f"Response: {availability}  |  Confidence: {row['confidence']['score']:.2f} ({row['confidence']['label']})").italic = True

        doc.add_paragraph(remarks)
        doc.add_paragraph(f"Sources: {', '.join(row['sources'])}")
        doc.add_paragraph("─" * 40)

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf


# ── Run ───────────────────────────────────────────────────────


# ── Feedback endpoints ────────────────────────────────────────


@app.post("/feedback/ingest")
def ingest_feedback_correction(data: dict, user: User = Depends(require("correct"))):
    """
    Ingest a human correction: (1) into Qdrant as a tier-1 golden answer,
    and (2) into the SHARED feedback log (feedback_log.jsonl) so every
    user's Feedback Loop page shows the same organisation-wide history —
    previously only Slack corrections reached the shared log, so web
    corrections were invisible to other users.
    """
    question    = str(data.get("question",    "") or "").strip()[:2000]
    good_answer = str(data.get("good_answer", "") or "").strip()[:5000]
    section     = str(data.get("section",     "") or "")[:200]
    raw_src     = str(data.get("source","feedback") or "").strip().lower()
    source      = raw_src if raw_src in {"workspace","assistant","review_queue","slack"} else "feedback"
    bad_answer  = str(data.get("bad_answer",  "") or "")[:5000]

    if not question or not good_answer:
        raise HTTPException(400, "question and good_answer are required")

    try:
        from ingest import ingest_correction
        vector_id = ingest_correction(question, good_answer, section, source)
    except Exception as e:
        raise _safe_error(e)

    # Persist to the shared org-wide ledger in Postgres so every user's
    # Feedback Loop page shows the same correction history.
    try:
        from sqlmodel import Session as _Session
        from database import engine as _engine
        from models import FeedbackPair as _FeedbackPair
        with _Session(_engine) as _s:
            _s.add(_FeedbackPair(
                signal="correction",
                question=question,
                good_answer=good_answer,
                bad_answer=bad_answer,
                section=section,
                source=source,
                user_name=user.name or user.email,
                user_email=user.email,
                confidence=float(data.get("confidence", 0) or 0),
            ))
            _s.commit()
    except Exception as e:
        print(f"  [feedback] postgres append failed: {e}")

    return {"status": "ingested", "vector_id": vector_id, "question": question[:80]}


@app.get("/feedback", dependencies=[Depends(require("feedback_read"))])
def get_feedback():
    """Return all feedback pairs from Postgres (shared org-wide ledger)."""
    from sqlmodel import Session as _Session, select as _select
    from database import engine as _engine
    from models import FeedbackPair as _FeedbackPair
    with _Session(_engine) as _s:
        rows = _s.exec(_select(_FeedbackPair).order_by(_FeedbackPair.created_at.desc())).all()
    pairs = [{
        "signal":      r.signal,
        "question":    r.question,
        "good_answer": r.good_answer,
        "bad_answer":  r.bad_answer,
        "section":     r.section,
        "source":      r.source,
        "user":        r.user_name,
        "email":       r.user_email,
        "confidence":  r.confidence,
        "logged_at":   r.created_at.isoformat() if r.created_at else None,
    } for r in rows]
    return {
        "pairs": pairs,
        "total": len(pairs),
        "thumbs_up":    sum(1 for p in pairs if p.get("signal") == "thumbs_up"),
        "thumbs_down":  sum(1 for p in pairs if p.get("signal") == "thumbs_down"),
        "corrections":  sum(1 for p in pairs if p.get("good_answer")),
    }




    log_path = Path(os.getenv("FEEDBACK_LOG", str(Path(__file__).parent / "feedback_log.jsonl")))
    if not log_path.exists():
        return PlainTextResponse("")

    lines = []
    for line in log_path.read_text().strip().splitlines():
        try:
            entry = json.loads(line)
            q = entry.get("question", "")
            a = entry.get("good_answer") or entry.get("bad_answer", "")
            if q and a:
                pair = {
                    "messages": [
                        {"role": "system", "content": "You are an enterprise RFP/RFI response assistant."},
                        {"role": "user",      "content": q},
                        {"role": "assistant", "content": a},
                    ]
                }
                lines.append(json.dumps(pair))
        except Exception:
            pass

    return PlainTextResponse(
        "\n".join(lines),
        media_type="application/jsonl",
        headers={"Content-Disposition": "attachment; filename=feedback.jsonl"}
    )


if __name__ == "__main__":
    import uvicorn
    uvicorn.run("api:app", host="0.0.0.0", port=8000, reload=True)