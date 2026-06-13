"""
parser.py — RFI/RFP document structure extractor

Parses uploaded documents and separates:
  - Section headers
  - Requirement statements
  - Questions
  - Evaluation criteria
  - Expected response fields

Supports: .xlsx, .xls, .docx, .pdf, .csv, .txt
Edge cases handled:
  - Merged cells in Excel (openpyxl returns None for merged slaves)
  - Multi-sheet workbooks (each sheet parsed independently)
  - Header rows anywhere in first 8 rows (not just row 1)
  - Numbered requirements like "1.", "1.1", "a)" without a "?"
  - Inline serial numbers stripped from question text
  - Duplicate questions deduplicated across sheets/pages
  - Empty/whitespace-only cells safely skipped
  - Non-UTF-8 encoded files (latin-1 fallback)
  - Password-protected Excel files (graceful skip)
  - Single-column sheets (treat every non-header row as a requirement)
  - PDF tables extracted in addition to text blocks
  - Very long cells (>2000 chars) treated as body text, not questions
  - Cells that are purely numeric (row numbers, scores) skipped
  - Google Sheets exported as .xlsx with hidden sheets skipped
  - DOCX tables where header row contains "requirement"/"description"
  - CSV files with BOM (UTF-8-sig), semicolon/tab delimiters
  - Files with no recognised questions → returns empty list gracefully
"""

import re
import io
import hashlib
from pathlib import Path
from datetime import datetime
from dataclasses import dataclass, asdict
from typing import Optional


# ── Data model ───────────────────────────────────────────────

@dataclass
class ExtractedItem:
    id: str
    section: str
    subsection: str
    question: str          # ONLY the question/requirement — no criteria text
    item_type: str         # "question" | "requirement" | "compliance" | "evaluation"
    priority: str          # "high" | "medium" | "low"
    source_row: Optional[int] = None
    raw_text: str = ""     # original text, for audit only — never echoed in answers


# ── Heuristics ───────────────────────────────────────────────

QUESTION_PATTERNS = [
    r'\?$',
    r'^(does|do|is|are|can|will|would|should|how|what|when|where|which|describe|explain|provide|list|specify|detail|mention|state|outline|elaborate)\b',
    r'^(please|kindly)\s+(describe|explain|provide|list|confirm|specify|share|outline)',
    r'^(support|provide|enable|allow|have|include|offer)\b',
    # Use-case / statement style
    r'^(user|system|platform|vendor|solution|tool|product)\s+(should|must|shall|can|needs? to|is able to)\b',
    # Action item style: "Configure SSO", "Enable MFA"
    r'^(configure|enable|disable|implement|verify|validate|ensure|review|assess|document|deploy|integrate|migrate|test|audit)\b',
    # Open action items: "Action:", "TODO:", "Open item:"
    r'^(action|todo|open item|pending|follow.?up|next step)\s*[:\-]',
]

SECTION_PATTERNS = [
    r'^\d+\.\s+[A-Z]',           # "1. SECTION NAME"
    r'^[A-Z][A-Z\s]{5,}$',       # "ALL CAPS HEADER"
    r'^(section|part|chapter)\s+\d+',
    r'^[IVX]+\.\s+[A-Z]',        # Roman numeral headers "IV. Security"
]

# Keywords that indicate a cell is ONLY scoring/criteria with no question content
CRITERIA_ONLY_PATTERNS = [
    r'^\s*(score|points?|marks?|weightage|max\s*score|total\s*marks?)[\s:=\d]',
    r'^\s*\d+\s*$',              # purely numeric (row numbers, scores)
    r'^\s*(yes|no|y|n|na|n\/a)\s*$',  # response cells
    r'^\s*(mandatory|optional|preferred)\s*$',  # standalone labels
]

# Column name keywords for auto-detection
QUESTION_COL_NAMES = [
    "requirement", "description", "question", "criteria",
    "item", "requirement description", "use case", "use-case",
    "feature", "capability", "parameter", "particulars", "details",
    "functionality", "specification", "spec", "query",
    "check", "checklist", "control", "assessment criteria",
    # Action item formats
    "action", "action item", "open action", "open item", "todo",
    "task", "activity", "deliverable", "open point",
    # Description formats
    "scope", "objective", "expectation", "need", "ask",
    "scenario", "test case", "test scenario", "acceptance criteria",
    "success criteria", "success criterion",
]
SECTION_COL_NAMES = [
    "category", "section", "area", "domain", "module",
    "sr", "s.no", "no.", "sl", "serial", "group", "phase",
    "workstream", "stream", "epic", "theme", "pillar", "track",
]
SUBSECTION_COL_NAMES = [
    "sub-category", "subcategory", "sub category", "sub-section",
    "subsection", "type", "sub-type", "sub type",
]
# Columns to explicitly skip — these are output/response columns
SKIP_COL_NAMES = [
    "response", "answer", "remarks", "comment", "vendor response",
    "matters response", "our response", "reply", "status",
    "compliance", "score", "weightage", "marks", "points",
    "evidence", "reference", "attachment",
    # Additional output columns
    "owner", "assignee", "assigned to", "responsible",
    "due date", "deadline", "target date", "eta",
    "priority level", "risk", "impact", "effort",
    "done", "completed", "closed", "resolution",
    "link", "url", "ticket", "jira",
]

HIGH_PRIORITY_KEYWORDS = [
    'mandatory', 'critical', 'required', 'must', 'security',
    'compliance', 'encryption', 'authentication', 'dlp', 'gdpr',
    'sebi', 'rbi', 'irdai', 'dpdp', 'pci', 'hipaa', 'soc2',
]

MAX_QUESTION_LENGTH = 1500   # cells longer than this are body text, not questions
MIN_QUESTION_LENGTH = 8      # cells shorter than this are skipped


# ── Helper utilities ─────────────────────────────────────────

def _safe_str(val) -> str:
    """Convert any cell value to a stripped string safely."""
    if val is None:
        return ""
    try:
        s = str(val).strip()
        # Remove leading serial numbers like "1.", "1.1.", "a)", "Q1."
        s = re.sub(r'^(\d+\.)+\s*', '', s)
        s = re.sub(r'^[a-zA-Z]\)\s*', '', s)
        s = re.sub(r'^Q\d+[.)]\s*', '', s, flags=re.IGNORECASE)
        return s.strip()
    except Exception:
        return ""


def _is_empty(text: str) -> bool:
    return not text or not text.strip()


def _is_criteria_only(text: str) -> bool:
    """True if the cell contains ONLY scoring/criteria info — no actual requirement."""
    t = text.lower().strip()
    if not t:
        return True
    return any(re.search(p, t) for p in CRITERIA_ONLY_PATTERNS)


def _is_skip_column(header: str) -> bool:
    """True if this column header indicates an output/response column to skip."""
    h = header.lower().strip()
    return any(k in h for k in SKIP_COL_NAMES)


def _is_section_header(text: str) -> bool:
    t = text.strip()
    if len(t) < 3 or len(t) > 120:
        return False
    return any(re.search(p, t, re.IGNORECASE) for p in SECTION_PATTERNS)


def _is_valid_question(text: str) -> bool:
    """
    Returns True if text looks like a requirement, question, use-case,
    action item, or description worth answering.

    Accepts:
      - Traditional questions ("Does the system support SSO?")
      - Statement requirements ("The platform must support MFA")
      - Use-cases ("Single Sign-On capability", "Data classification for PII")
      - Action items ("Configure LDAP integration", "Verify encryption at rest")
      - Open action items ("Action: Review DLP policy")
      - Descriptions (any noun phrase 3+ words, 8–1500 chars)

    Rejects:
      - Pure numbers, single words, URLs
      - Purely scoring/criteria cells
      - Section headers (checked separately before this)
      - Cells over 1500 chars (body paragraphs, not discrete requirements)
    """
    t = text.strip()
    if len(t) < MIN_QUESTION_LENGTH or len(t) > MAX_QUESTION_LENGTH:
        return False
    if re.match(r'^\d+(\.\d+)?$', t):      # purely numeric
        return False
    if re.match(r'^https?://', t):           # URLs
        return False
    if _is_criteria_only(t):
        return False
    words = t.split()
    if len(words) < 2:                      # single-word cells
        return False
    # Accept anything with 3+ words — covers noun-phrase use-cases and descriptions
    # For 2-word cells, require a known question/action pattern
    if len(words) == 2:
        t_lower = t.lower()
        if not any(re.search(p, t_lower, re.IGNORECASE) for p in QUESTION_PATTERNS):
            return False
    return True


def _clean_question(text: str) -> str:
    """
    Strip criteria/scoring text appended to a question.
    e.g. "Does the system support SSO? (Mandatory, 10 points)" → "Does the system support SSO?"
    """
    if not text:
        return ""
    # Remove parenthetical criteria
    text = re.sub(r'\s*\([^)]*(?:mandatory|points?|marks?|score|criteria|required|optional)[^)]*\)', '', text, flags=re.IGNORECASE)
    # Remove bracket criteria [Mandatory]
    text = re.sub(r'\s*\[[^\]]*(?:mandatory|points?|marks?|score|required)[^\]]*\]', '', text, flags=re.IGNORECASE)
    # Remove trailing scores "- 10 marks", "– 5 points"
    text = re.sub(r'\s*[-–]\s*\d+\s*(marks?|points?|%)\s*$', '', text, flags=re.IGNORECASE)
    # Remove "Mandatory:" / "Required:" prefixes
    text = re.sub(r'^(mandatory|required|optional|critical)\s*:\s*', '', text, flags=re.IGNORECASE)
    # Remove leading serial numbers added again after previous clean
    text = re.sub(r'^(\d+[.)]\s*)+', '', text)
    # Collapse multiple spaces
    text = re.sub(r'\s{2,}', ' ', text)
    return text.strip()


def _classify_type(text: str) -> str:
    t = text.lower().strip()
    if re.search(r'\?$', t) or re.search(r'^(does|do|is|are|can|will)\b', t):
        return "question"
    if re.search(r'(mandatory|shall|must|compliance|regulatory|regulation)', t):
        return "compliance"
    if re.search(r'(score|points?|marks?|weightage|evaluation criteria)', t):
        return "evaluation"
    if re.search(r'^(action|todo|open item|pending|follow.?up|next step)\s*[:\-]', t):
        return "action_item"
    if re.search(r'^(configure|enable|disable|implement|verify|validate|ensure|review|deploy|integrate|test|audit)\b', t):
        return "action_item"
    if re.search(r'(use case|use-case|scenario|user story)', t):
        return "use_case"
    return "requirement"


def _classify_priority(text: str) -> str:
    t = text.lower()
    if any(k in t for k in HIGH_PRIORITY_KEYWORDS):
        return "high"
    if re.search(r'(should|preferred|desirable|optional|nice.to.have)', t):
        return "low"
    return "medium"


def _dedup_key(text: str) -> str:
    """Generate a deduplication key — normalised lowercase fingerprint."""
    normalised = re.sub(r'\W+', '', text.lower())
    return hashlib.md5(normalised.encode()).hexdigest()


# ── Excel parser ─────────────────────────────────────────────

def _detect_header_row(sheet) -> tuple:
    """
    Scan the first 8 rows to find a header row.
    Returns (header_row_index, col_mapping_dict) or (None, {}).

    col_mapping_dict keys:
      question_col  → int index of the requirements column
      section_col   → int index of category/section column (optional)
      subsection_col → int index of sub-category column (optional)
      skip_cols     → set of int indices to ignore (response/score columns)
    """
    for row_idx, row in enumerate(sheet.iter_rows(max_row=8, values_only=True), start=1):
        cells = [str(c).strip().lower() if c is not None else "" for c in row]
        non_empty = [c for c in cells if c]
        if len(non_empty) < 1:
            continue

        q_match  = sum(1 for c in cells if any(k in c for k in QUESTION_COL_NAMES))
        s_match  = sum(1 for c in cells if any(k in c for k in SECTION_COL_NAMES))

        if q_match == 0 and s_match == 0:
            continue

        mapping = {"skip_cols": set()}
        for col_idx, cell in enumerate(cells):
            if not cell:
                continue
            if _is_skip_column(cell):
                mapping["skip_cols"].add(col_idx)
                continue
            if any(k in cell for k in QUESTION_COL_NAMES) and "question_col" not in mapping:
                mapping["question_col"] = col_idx
            elif any(k in cell for k in SECTION_COL_NAMES) and "section_col" not in mapping:
                mapping["section_col"] = col_idx
            elif any(k in cell for k in SUBSECTION_COL_NAMES) and "subsection_col" not in mapping:
                mapping["subsection_col"] = col_idx

        if "question_col" in mapping:
            return row_idx, mapping

    return None, {}


def _parse_excel(buf: io.BytesIO, filename: str) -> list[ExtractedItem]:
    try:
        import openpyxl
    except ImportError:
        raise ImportError("pip install openpyxl")

    try:
        wb = openpyxl.load_workbook(buf, data_only=True)
    except Exception as e:
        # Password-protected or corrupt file
        raise ValueError(f"Cannot open Excel file (may be password-protected): {e}")

    items = []
    item_counter = 0
    seen: set = set()   # deduplication across sheets

    for sheet in wb.worksheets:
        # Skip hidden sheets (common in Google Sheets exports)
        if sheet.sheet_state == "hidden":
            continue

        # Skip sheets with fewer than 2 rows
        if sheet.max_row < 2:
            continue

        header_row, col_map = _detect_header_row(sheet)

        if header_row and col_map:
            # ── Structured table mode ──────────────────────────
            current_section    = sheet.title
            current_subsection = ""
            q_col   = col_map.get("question_col", 0)
            s_col   = col_map.get("section_col")
            ss_col  = col_map.get("subsection_col")
            skip    = col_map.get("skip_cols", set())

            for row_idx, row in enumerate(sheet.iter_rows(values_only=True), start=1):
                if row_idx <= header_row:
                    continue

                cells = [_safe_str(c) for c in row]

                # Skip fully empty rows
                if not any(c for c in cells):
                    continue

                # Update section/subsection from their columns (carry-forward)
                if s_col is not None and s_col < len(cells) and cells[s_col]:
                    current_section = cells[s_col]
                if ss_col is not None and ss_col < len(cells) and cells[ss_col]:
                    current_subsection = cells[ss_col]

                # Get question text
                question_text = cells[q_col] if q_col < len(cells) else ""

                # If question col is empty but other non-skip cols have content,
                # try the longest non-skip cell (handles misaligned columns)
                if not question_text or len(question_text) < MIN_QUESTION_LENGTH:
                    candidates = [
                        cells[i] for i in range(len(cells))
                        if i not in skip
                        and i != s_col
                        and i != ss_col
                        and len(cells[i]) >= MIN_QUESTION_LENGTH
                    ]
                    if candidates:
                        question_text = max(candidates, key=len)

                if not _is_valid_question(question_text):
                    continue

                cleaned = _clean_question(question_text)
                if not _is_valid_question(cleaned):
                    continue

                # Deduplicate
                dk = _dedup_key(cleaned)
                if dk in seen:
                    continue
                seen.add(dk)

                item_counter += 1
                items.append(ExtractedItem(
                    id=f"Q{item_counter:03d}",
                    section=current_section or sheet.title,
                    subsection=current_subsection,
                    question=cleaned,
                    item_type=_classify_type(cleaned),
                    priority=_classify_priority(cleaned),
                    source_row=row_idx,
                    raw_text=question_text,
                ))

        else:
            # ── Fallback: single-column or unstructured sheet ──
            current_section = sheet.title

            # If sheet has only 1 non-empty column, treat every row as a requirement
            all_rows = list(sheet.iter_rows(values_only=True))
            col_counts = [sum(1 for r in all_rows if r[i] is not None) for i in range(sheet.max_column or 1)]
            is_single_col = sum(1 for c in col_counts if c > 0) <= 2

            for row_idx, row in enumerate(all_rows, start=1):
                cells = [_safe_str(c) for c in row]
                non_empty = [c for c in cells if c]
                if not non_empty:
                    continue

                first = non_empty[0]

                if _is_section_header(first) and not is_single_col:
                    current_section = first
                    continue

                # Pick the best candidate cell
                if is_single_col:
                    question_text = first
                else:
                    question_text = ""
                    for c in cells:
                        if _is_valid_question(c):
                            question_text = c
                            break

                if not _is_valid_question(question_text):
                    continue

                cleaned = _clean_question(question_text)
                if not _is_valid_question(cleaned):
                    continue

                dk = _dedup_key(cleaned)
                if dk in seen:
                    continue
                seen.add(dk)

                item_counter += 1
                items.append(ExtractedItem(
                    id=f"Q{item_counter:03d}",
                    section=current_section,
                    subsection="",
                    question=cleaned,
                    item_type=_classify_type(cleaned),
                    priority=_classify_priority(cleaned),
                    source_row=row_idx,
                    raw_text=question_text,
                ))

    return items


# ── DOCX parser ───────────────────────────────────────────────

def _parse_docx(buf: io.BytesIO) -> list[ExtractedItem]:
    try:
        from docx import Document
    except ImportError:
        raise ImportError("pip install python-docx")

    try:
        doc = Document(buf)
    except Exception as e:
        raise ValueError(f"Cannot open Word document: {e}")

    items = []
    item_counter = 0
    seen: set = set()
    current_section    = "General"
    current_subsection = ""

    # ── Paragraphs ───────────────────────────────────────────
    for para in doc.paragraphs:
        text = para.text.strip()
        if _is_empty(text):
            continue

        # Heading detection
        if para.style and para.style.name.startswith("Heading"):
            try:
                level = int(para.style.name.split()[-1])
            except (ValueError, IndexError):
                level = 1
            if level == 1:
                current_section    = text
                current_subsection = ""
            else:
                current_subsection = text
            continue

        if _is_section_header(text):
            current_section = text
            continue

        if not _is_valid_question(text):
            continue

        cleaned = _clean_question(text)
        if not _is_valid_question(cleaned):
            continue

        dk = _dedup_key(cleaned)
        if dk in seen:
            continue
        seen.add(dk)

        item_counter += 1
        items.append(ExtractedItem(
            id=f"Q{item_counter:03d}",
            section=current_section,
            subsection=current_subsection,
            question=cleaned,
            item_type=_classify_type(cleaned),
            priority=_classify_priority(cleaned),
            raw_text=text,
        ))

    # ── Tables inside DOCX ───────────────────────────────────
    for table in doc.tables:
        if not table.rows:
            continue

        # Try to detect if the first row is a header
        first_row_cells = [c.text.strip().lower() for c in table.rows[0].cells]
        has_header = any(
            any(k in c for k in QUESTION_COL_NAMES + SECTION_COL_NAMES)
            for c in first_row_cells
        )

        q_col_idx  = None
        s_col_idx  = None
        ss_col_idx = None
        skip_cols: set = set()

        if has_header:
            for i, c in enumerate(first_row_cells):
                if _is_skip_column(c):
                    skip_cols.add(i)
                elif any(k in c for k in QUESTION_COL_NAMES) and q_col_idx is None:
                    q_col_idx = i
                elif any(k in c for k in SECTION_COL_NAMES) and s_col_idx is None:
                    s_col_idx = i
                elif any(k in c for k in SUBSECTION_COL_NAMES) and ss_col_idx is None:
                    ss_col_idx = i

        start_row = 1 if has_header else 0
        table_section    = current_section
        table_subsection = current_subsection

        for row in table.rows[start_row:]:
            cell_texts = [_safe_str(c.text) for c in row.cells]

            if q_col_idx is not None:
                question_text = cell_texts[q_col_idx] if q_col_idx < len(cell_texts) else ""
                if s_col_idx is not None and s_col_idx < len(cell_texts) and cell_texts[s_col_idx]:
                    table_section = cell_texts[s_col_idx]
                if ss_col_idx is not None and ss_col_idx < len(cell_texts) and cell_texts[ss_col_idx]:
                    table_subsection = cell_texts[ss_col_idx]
            else:
                # Pick longest valid cell
                candidates = [t for i, t in enumerate(cell_texts) if i not in skip_cols and _is_valid_question(t)]
                question_text = max(candidates, key=len) if candidates else ""

            if not _is_valid_question(question_text):
                continue

            cleaned = _clean_question(question_text)
            if not _is_valid_question(cleaned):
                continue

            dk = _dedup_key(cleaned)
            if dk in seen:
                continue
            seen.add(dk)

            item_counter += 1
            items.append(ExtractedItem(
                id=f"Q{item_counter:03d}",
                section=table_section,
                subsection=table_subsection,
                question=cleaned,
                item_type=_classify_type(cleaned),
                priority=_classify_priority(cleaned),
                raw_text=question_text,
            ))

    return items


# ── PDF parser ────────────────────────────────────────────────

def _parse_pdf(buf: io.BytesIO) -> list[ExtractedItem]:
    try:
        import pdfplumber
    except ImportError:
        raise ImportError("pip install pdfplumber")

    items = []
    item_counter = 0
    seen: set = set()
    current_section = "General"

    try:
        with pdfplumber.open(buf) as pdf:
            for page in pdf.pages:

                # ── Extract tables first ──────────────────────
                for table in (page.extract_tables() or []):
                    if not table or len(table) < 2:
                        continue

                    # Detect header row
                    header = [str(c).strip().lower() if c else "" for c in table[0]]
                    has_header = any(
                        any(k in c for k in QUESTION_COL_NAMES)
                        for c in header
                    )
                    q_col = None
                    skip_cols: set = set()

                    if has_header:
                        for i, c in enumerate(header):
                            if _is_skip_column(c):
                                skip_cols.add(i)
                            elif any(k in c for k in QUESTION_COL_NAMES) and q_col is None:
                                q_col = i

                    start = 1 if has_header else 0
                    for row in table[start:]:
                        if not row:
                            continue
                        cells = [_safe_str(c) for c in row]

                        if q_col is not None and q_col < len(cells):
                            question_text = cells[q_col]
                        else:
                            candidates = [t for i, t in enumerate(cells) if i not in skip_cols and _is_valid_question(t)]
                            question_text = max(candidates, key=len) if candidates else ""

                        if not _is_valid_question(question_text):
                            continue
                        cleaned = _clean_question(question_text)
                        if not _is_valid_question(cleaned):
                            continue

                        dk = _dedup_key(cleaned)
                        if dk in seen:
                            continue
                        seen.add(dk)

                        item_counter += 1
                        items.append(ExtractedItem(
                            id=f"Q{item_counter:03d}",
                            section=current_section,
                            subsection="",
                            question=cleaned,
                            item_type=_classify_type(cleaned),
                            priority=_classify_priority(cleaned),
                            raw_text=question_text,
                        ))

                # ── Extract body text ─────────────────────────
                text = page.extract_text() or ""
                for line in text.split("\n"):
                    line = _safe_str(line)
                    if _is_empty(line):
                        continue

                    if _is_section_header(line):
                        current_section = line
                        continue

                    if not _is_valid_question(line):
                        continue

                    cleaned = _clean_question(line)
                    if not _is_valid_question(cleaned):
                        continue

                    dk = _dedup_key(cleaned)
                    if dk in seen:
                        continue
                    seen.add(dk)

                    item_counter += 1
                    items.append(ExtractedItem(
                        id=f"Q{item_counter:03d}",
                        section=current_section,
                        subsection="",
                        question=cleaned,
                        item_type=_classify_type(cleaned),
                        priority=_classify_priority(cleaned),
                        raw_text=line,
                    ))

    except Exception as e:
        raise ValueError(f"Cannot parse PDF: {e}")

    return items


# ── CSV parser ────────────────────────────────────────────────

def _parse_csv(buf: io.BytesIO) -> list[ExtractedItem]:
    import csv

    buf.seek(0)
    raw = buf.read()

    # Try UTF-8-sig (BOM), then UTF-8, then latin-1
    for encoding in ("utf-8-sig", "utf-8", "latin-1"):
        try:
            text = raw.decode(encoding)
            break
        except UnicodeDecodeError:
            continue
    else:
        text = raw.decode("latin-1", errors="replace")

    # Detect delimiter: comma, semicolon, tab, pipe
    sample = text[:2000]
    delimiter = ","
    for d in ["\t", ";", "|"]:
        if text.count(d) > text.count(","):
            delimiter = d
            break

    reader = csv.DictReader(io.StringIO(text), delimiter=delimiter)
    items = []
    item_counter = 0
    seen: set = set()
    current_section = "General"

    # Find question column
    question_col = None
    section_col  = None
    skip_set: set = set()

    if reader.fieldnames:
        for fn in reader.fieldnames:
            if not fn:
                continue
            fn_lower = fn.lower().strip()
            if _is_skip_column(fn_lower):
                skip_set.add(fn)
            elif any(k in fn_lower for k in QUESTION_COL_NAMES) and question_col is None:
                question_col = fn
            elif any(k in fn_lower for k in SECTION_COL_NAMES) and section_col is None:
                section_col = fn

        # Fallback: use first non-skip column
        if question_col is None:
            for fn in reader.fieldnames:
                if fn and fn not in skip_set:
                    question_col = fn
                    break

    for row in reader:
        if section_col and row.get(section_col, "").strip():
            current_section = row[section_col].strip()

        text_val = row.get(question_col, "") if question_col else ""
        text_val = _safe_str(text_val)

        if not _is_valid_question(text_val):
            continue

        cleaned = _clean_question(text_val)
        if not _is_valid_question(cleaned):
            continue

        dk = _dedup_key(cleaned)
        if dk in seen:
            continue
        seen.add(dk)

        item_counter += 1
        items.append(ExtractedItem(
            id=f"Q{item_counter:03d}",
            section=current_section,
            subsection="",
            question=cleaned,
            item_type=_classify_type(cleaned),
            priority=_classify_priority(cleaned),
            raw_text=text_val,
        ))

    return items


# ── TXT parser ────────────────────────────────────────────────

def _parse_txt(buf: io.BytesIO) -> list[ExtractedItem]:
    buf.seek(0)
    raw = buf.read()
    for encoding in ("utf-8-sig", "utf-8", "latin-1"):
        try:
            content = raw.decode(encoding)
            break
        except UnicodeDecodeError:
            continue
    else:
        content = raw.decode("latin-1", errors="replace")

    items = []
    item_counter = 0
    seen: set = set()
    current_section = "General"

    for line in content.split("\n"):
        line = _safe_str(line)
        if _is_empty(line):
            continue

        if _is_section_header(line):
            current_section = line
            continue

        if not _is_valid_question(line):
            continue

        cleaned = _clean_question(line)
        if not _is_valid_question(cleaned):
            continue

        dk = _dedup_key(cleaned)
        if dk in seen:
            continue
        seen.add(dk)

        item_counter += 1
        items.append(ExtractedItem(
            id=f"Q{item_counter:03d}",
            section=current_section,
            subsection="",
            question=cleaned,
            item_type=_classify_type(cleaned),
            priority=_classify_priority(cleaned),
            raw_text=line,
        ))

    return items


# ── Main entry point ─────────────────────────────────────────

def parse_rfp_document(buf: io.BytesIO, filename: str) -> list[ExtractedItem]:
    """
    Parse an RFP/RFI document and return a list of clean ExtractedItems.

    Each item contains ONLY the question/requirement — no criteria text,
    no scoring info, no response columns.

    Returns an empty list (never raises) if no questions are found.
    """
    ext = Path(filename).suffix.lower()

    # Strip query strings from filenames like "file.xlsx?dl=1"
    if "?" in ext:
        ext = ext.split("?")[0]

    buf.seek(0)

    try:
        if ext in (".xlsx", ".xls"):
            return _parse_excel(buf, filename)
        elif ext == ".docx":
            return _parse_docx(buf)
        elif ext == ".pdf":
            return _parse_pdf(buf)
        elif ext == ".csv":
            return _parse_csv(buf)
        elif ext in (".txt", ".md"):
            return _parse_txt(buf)
        else:
            # Unknown extension — try Excel first, then plain text
            try:
                buf.seek(0)
                result = _parse_excel(buf, filename)
                if result:
                    return result
            except Exception:
                pass
            buf.seek(0)
            return _parse_txt(buf)
    except Exception as e:
        # Never crash the API — return empty list with a warning
        import sys
        print(f"[parser] Warning: failed to parse {filename}: {e}", file=sys.stderr)
        return []


def items_to_dict(items: list[ExtractedItem]) -> list[dict]:
    return [asdict(i) for i in items]


# ── CLI test ─────────────────────────────────────────────────

if __name__ == "__main__":
    import sys
    import json

    if len(sys.argv) < 2:
        print("Usage: python3 parser.py <path_to_rfp_file> [--json]")
        sys.exit(1)

    path = Path(sys.argv[1])
    if not path.exists():
        print(f"File not found: {path}")
        sys.exit(1)

    with open(path, "rb") as f:
        buf = io.BytesIO(f.read())

    items = parse_rfp_document(buf, path.name)
    print(f"\n✅ Extracted {len(items)} questions/requirements from {path.name}\n")

    for item in items[:10]:
        print(f"  [{item.id}] [{item.section}] [{item.priority.upper()}] [{item.item_type}]")
        print(f"       {item.question[:120]}")
        print()

    if len(items) > 10:
        print(f"  ... and {len(items) - 10} more\n")

    if "--json" in sys.argv:
        print(json.dumps(items_to_dict(items), indent=2))