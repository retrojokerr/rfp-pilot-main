"""
ingest.py — Document ingestion pipeline (updated)

Changes from original:
  - Stores upload_date, version, source_type metadata in every vector
  - Semantic chunking preserves section context
  - Supports both local files and Google Drive buffers
  - Skip-if-unchanged: files not modified since last sync are skipped
"""

import os
import io
import hashlib
from pathlib import Path
from datetime import datetime, timezone
from dotenv import load_dotenv

from sentence_transformers import SentenceTransformer
from qdrant_client import QdrantClient
from qdrant_client.models import (
    Distance, VectorParams, PointStruct,
    Filter, FieldCondition, MatchValue,
)

load_dotenv()

# ── Init ─────────────────────────────────────────────────────

embedder = SentenceTransformer("BAAI/bge-small-en-v1.5")

qdrant = QdrantClient(
    url=os.getenv("QDRANT_URL", "http://localhost:6333"),
    api_key=os.getenv("QDRANT_API_KEY"),
)

COLLECTION    = "rfi_knowledge_base"
VECTOR_SIZE   = 384
CHUNK_SIZE    = 400
CHUNK_OVERLAP = 80


# ── Collection setup ──────────────────────────────────────────

def ensure_collection():
    existing = [c.name for c in qdrant.get_collections().collections]
    if COLLECTION not in existing:
        qdrant.create_collection(
            collection_name=COLLECTION,
            vectors_config=VectorParams(size=VECTOR_SIZE, distance=Distance.COSINE),
        )
        print(f"Collection '{COLLECTION}' created.")
    else:
        print(f"Collection '{COLLECTION}' already exists.")


# ── Text extraction ───────────────────────────────────────────

def extract_text_from_buffer(buf: io.BytesIO, filename: str) -> list[dict]:
    ext = Path(filename).suffix.lower()
    buf.seek(0)
    if ext in (".xlsx", ".xls"):
        return _extract_excel(buf)
    elif ext == ".docx":
        return _extract_docx(buf)
    elif ext == ".pdf":
        return _extract_pdf(buf)
    elif ext == ".csv":
        return _extract_csv(buf)
    elif ext in (".txt", ".md"):
        text = buf.read().decode("utf-8", errors="replace")
        return _chunk_text(text)
    else:
        raise ValueError(f"Unsupported file type: {ext}")


def _chunk_text(text: str, section: str = "") -> list[dict]:
    words = text.split()
    chunks = []
    i = 0
    while i < len(words):
        chunk_words = words[i:i + CHUNK_SIZE]
        chunk_text  = " ".join(chunk_words).strip()
        if len(chunk_text) > 50:
            chunks.append({"text": chunk_text, "section": section, "page": 0})
        i += CHUNK_SIZE - CHUNK_OVERLAP
    return chunks


def _extract_excel(buf: io.BytesIO) -> list[dict]:
    import openpyxl
    wb = openpyxl.load_workbook(buf, data_only=True)
    chunks = []
    for sheet in wb.worksheets:
        rows_text = []
        for row in sheet.iter_rows(values_only=True):
            cells = [str(c).strip() for c in row if c is not None and str(c).strip()]
            if cells:
                rows_text.append(" | ".join(cells))
        full_text = "\n".join(rows_text)
        chunks.extend(_chunk_text(full_text, section=sheet.title))
    return chunks


def _extract_docx(buf: io.BytesIO) -> list[dict]:
    from docx import Document
    doc = Document(buf)
    current_section = ""
    chunks = []
    block = []
    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            continue
        if para.style.name.startswith("Heading"):
            if block:
                chunks.extend(_chunk_text(" ".join(block), section=current_section))
                block = []
            current_section = text
        else:
            block.append(text)
    if block:
        chunks.extend(_chunk_text(" ".join(block), section=current_section))
    for table in doc.tables:
        rows = []
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if cells:
                rows.append(" | ".join(cells))
        if rows:
            chunks.extend(_chunk_text("\n".join(rows), section="Table"))
    return chunks


def _extract_pdf(buf: io.BytesIO) -> list[dict]:
    import pdfplumber
    chunks = []
    with pdfplumber.open(buf) as pdf:
        for page_num, page in enumerate(pdf.pages, start=1):
            text = page.extract_text() or ""
            page_chunks = _chunk_text(text)
            for c in page_chunks:
                c["page"] = page_num
            chunks.extend(page_chunks)
    return chunks


def _extract_csv(buf: io.BytesIO) -> list[dict]:
    import csv
    buf.seek(0)
    text = buf.read().decode("utf-8", errors="replace")
    rows = list(csv.reader(io.StringIO(text)))
    lines = [" | ".join(r) for r in rows if any(c.strip() for c in r)]
    return _chunk_text("\n".join(lines))


# ── Ingestion ──────────────────────────────────────────────────

def ingest_file(
    buf: io.BytesIO,
    filename: str,
    file_id: str = "",
    upload_date: datetime = None,
    version: str = "1.0",
    source_type: str = "manual",
    tags: list[str] = None,
):
    ensure_collection()
    upload_date = upload_date or datetime.now(timezone.utc)
    upload_date_str = upload_date.isoformat()
    tags = tags or []

    print(f"  Extracting text from {filename}...")
    try:
        text_chunks = extract_text_from_buffer(buf, filename)
    except Exception as e:
        print(f"  ✗ Failed to extract {filename}: {e}")
        return

    if not text_chunks:
        print(f"  ✗ No text extracted from {filename}")
        return

    # Remove old vectors for this file before re-ingesting
    try:
        qdrant.delete(
            collection_name=COLLECTION,
            points_selector=Filter(
                must=[FieldCondition(key="source_file", match=MatchValue(value=filename))]
            ),
        )
    except Exception:
        pass

    texts   = [c["text"] for c in text_chunks]
    vectors = embedder.encode(texts, normalize_embeddings=True, batch_size=32).tolist()

    points = []
    for i, (chunk, vector) in enumerate(zip(text_chunks, vectors)):
        point_id = int(hashlib.md5(f"{filename}_{i}".encode()).hexdigest()[:8], 16)
        points.append(PointStruct(
            id=point_id,
            vector=vector,
            payload={
                "text":        chunk["text"],
                "source_file": filename,
                "file_id":     file_id,
                "section":     chunk.get("section", ""),
                "page":        chunk.get("page", 0),
                "upload_date": upload_date_str,
                "version":     version,
                "source_type": source_type,
                "tags":        tags,
            },
        ))

    qdrant.upsert(collection_name=COLLECTION, points=points)
    print(f"  ✓ {filename}: {len(points)} vectors ingested (upload_date: {upload_date_str[:10]})")


# ── Google Drive ingestion ────────────────────────────────────

GOOGLE_MIME_EXPORT = {
    "application/vnd.google-apps.document":     ("application/pdf", ".pdf"),
    "application/vnd.google-apps.spreadsheet":  ("application/vnd.openxmlformats-officedocument.spreadsheetml.sheet", ".xlsx"),
    "application/vnd.google-apps.presentation": ("application/pdf", ".pdf"),
}


def download_file(service, file_id: str, filename: str, mime_type: str = None):
    from googleapiclient.http import MediaIoBaseDownload
    if mime_type in GOOGLE_MIME_EXPORT:
        export_mime, new_ext = GOOGLE_MIME_EXPORT[mime_type]
        request  = service.files().export_media(fileId=file_id, mimeType=export_mime)
        filename = Path(filename).stem + new_ext
    else:
        request = service.files().get_media(fileId=file_id, supportsAllDrives=True)
    buf = io.BytesIO()
    downloader = MediaIoBaseDownload(buf, request)
    done = False
    while not done:
        _, done = downloader.next_chunk()
    buf.seek(0)
    return buf, filename


def get_ingested_modified_time(file_id: str) -> str | None:
    """
    Return the upload_date stored in Qdrant for the given Drive file_id,
    or None if the file has never been ingested.
    Used by run_ingestion to skip files unchanged since last sync.
    """
    try:
        results = qdrant.scroll(
            collection_name=COLLECTION,
            scroll_filter=Filter(
                must=[FieldCondition(key="file_id", match=MatchValue(value=file_id))]
            ),
            limit=1,
            with_payload=True,
            with_vectors=False,
        )
        points = results[0]
        if points:
            return points[0].payload.get("upload_date")
    except Exception:
        pass
    return None


def run_ingestion():
    """Run full ingestion from Google Drive, skipping unchanged files."""
    from google.oauth2.service_account import Credentials
    from googleapiclient.discovery import build

    creds   = Credentials.from_service_account_file(
        "credentials.json",
        scopes=["https://www.googleapis.com/auth/drive.readonly"]
    )
    service = build("drive", "v3", credentials=creds)

    folder_id = os.getenv("DRIVE_FOLDER_ID")
    results = service.files().list(
        q=f"'{folder_id}' in parents and trashed=false",
        fields="files(id,name,mimeType,modifiedTime)",
        pageSize=200,
        includeItemsFromAllDrives=True,
        supportsAllDrives=True,
    ).execute()

    files = results.get("files", [])
    print(f"\nFound {len(files)} files in Drive folder\n")

    skipped = 0
    processed = 0

    for f in files:
        name      = f["name"]
        mime_type = f["mimeType"]
        ext       = Path(name).suffix.lower()

        is_google = mime_type in GOOGLE_MIME_EXPORT
        is_ok_ext = ext in (".pdf", ".docx", ".xlsx", ".csv", ".txt")

        if not is_google and not is_ok_ext:
            print(f"  Skipping {name} (unsupported type)")
            continue

        modified = f.get("modifiedTime", "")

        # Skip if file hasn't changed since last sync
        last_ingested = get_ingested_modified_time(f["id"])
        if last_ingested and modified:
            # Parse as real datetimes — string comparison fails on
            # millisecond format differences (.000Z vs +00:00)
            try:
                mod_dt  = datetime.fromisoformat(modified.replace("Z", "+00:00"))
                last_dt = datetime.fromisoformat(last_ingested.replace("Z", "+00:00"))
                if mod_dt <= last_dt:
                    print(f"  ↩ Skipping {name} (unchanged since {last_dt.date()})")
                    skipped += 1
                    continue
            except (ValueError, TypeError):
                pass  # can't parse — re-ingest to be safe

        print(f"\nProcessing: {name} (modified: {modified[:10] if modified else 'unknown'})")
        buf, name = download_file(service, f["id"], name, mime_type)

        upload_date = (
            datetime.fromisoformat(modified.replace("Z", "+00:00"))
            if modified else None
        )
        ingest_file(
            buf=buf,
            filename=name,
            file_id=f["id"],
            upload_date=upload_date,
            source_type="drive",
        )
        processed += 1

    print(f"\n✅ Ingestion complete — {processed} processed, {skipped} skipped (unchanged)")


if __name__ == "__main__":
    run_ingestion()


# ── Feedback correction ingestion ─────────────────────────────

KNOWN_CORRECTION_SOURCES = {"slack", "assistant", "review_queue", "workspace", "feedback"}


def ingest_correction(question: str, good_answer: str, section: str = "", source: str = "feedback") -> str:
    """
    Ingest a human-verified correction as a tier-1 golden answer vector.
    Last-correction-wins: same question always upserts over the previous answer.
    Source is recorded for provenance (slack | assistant | review_queue | workspace).
    """
    ensure_collection()

    source = (source or "feedback").strip().lower()
    if source not in KNOWN_CORRECTION_SOURCES:
        source = "feedback"

    now = datetime.now(timezone.utc).isoformat()
    combined_text = f"Q: {question}\nA: {good_answer}"

    vector = embedder.encode([question], normalize_embeddings=True).tolist()[0]

    norm_q   = " ".join(question.lower().split())
    point_id = int(hashlib.md5(f"correction_{norm_q}".encode()).hexdigest()[:8], 16)

    qdrant.upsert(
        collection_name=COLLECTION,
        points=[PointStruct(
            id=point_id,
            vector=vector,
            payload={
                "text":              combined_text,
                "source_file":       f"[Correction] {source}",
                "file_id":           f"correction_{point_id}",
                "section":           section,
                "page":              0,
                "upload_date":       now,
                "version":           1,
                "source_type":       "golden_answer",
                "tags":              ["correction", "feedback", source],
                "is_correction":     True,
                "original_question": question,
                "correction_source": source,
                "corrected_at":      now,
            }
        )]
    )

    print(f"  Correction ingested (source={source}): {question[:60]}...")
    return str(point_id)